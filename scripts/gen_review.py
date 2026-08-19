#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
gen_review.py — 面试复盘 Word 文档生成器（interview-review skill 核心脚本）

把一份结构化的复盘数据(JSON) 渲染成格式统一的 Word 文档。
支持「新建」首份文档 和 「追加」新面试场次到已有文档两种模式。

用法:
  # 首次创建：从一个面试场次数据生成复盘文档
  python gen_review.py --session session.json --out 复盘.docx

  # 追加：把新场次追加到已有文档（不重建封面/说明，只追加该场章节）
  python gen_review.py --session session.json --out 已有复盘.docx --append

session.json 结构见 references/session-schema.md，核心字段:
{
  "candidate": "姓名",
  "round": 1,
  "company": "公司",
  "position": "岗位",
  "date": "2026-07-29",
  "result": "未通过(一面)",
  "jd": { "responsibilities": [...], "requirements": [...] },     # 可选
  "business_line": "面试官所在业务线描述",                          # 可选
  "match_hint": "匹配度提示",                                      # 可选
  "qa": [
    { "no":"Q1", "time":"00:03", "module":"模块",
      "q":"面试官问题(原话)", "a":"我的回答(原话)",
      "eval":"评价 ✅🟡🔴 + 点评", "better":"改进版回答", "gain":"提升点" }
  ],
  "interviewer_feedback": "面试官原话反馈",                        # 可选
  "conclusions": [ ["标签","结论"], ... ]                          # 可选
}
"""
import argparse
import json
import os
import sys

try:
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    sys.stderr.write("[gen_review] python-docx required: pip install python-docx\n")
    raise

# 主题色（品牌蓝系，可按公司调整）
COLORS = {
    "primary": "2932E1", "accent": "4E6EF2", "muted": "8A8F99",
    # 普通题（面试官问、你答）
    "q_bg": "EDF0FF", "a_bg": "FFFFFF", "eval_bg": "FFF4E6",
    "better_bg": "EAF7EE", "gain_bg": "FDEAEA",
    "q_text": "2932E1", "a_text": "2C3E50", "eval_text": "8A6D3B",
    "better_text": "1E7E34", "gain_text": "C0392B",
    "label_text": "555555", "label_bg": "F5F6FA",
    # 反问环节（你问、面试官答）——用紫色系区分，标签也标明"我的提问/面试官回答"
    "myq_bg": "F3EDFF", "iva_bg": "FFFFFF",
    "myq_text": "6F42C1", "iva_text": "2C3E50",
    # 评价两侧：提问评价用橙系(同eval)，回答分析用青系
    "myq_eval_bg": "FFF4E6", "myq_eval_text": "8A6D3B",
    "myq_better_bg": "EAF7EE", "myq_better_text": "1E7E34",
    "myq_gain_bg": "FDEAEA", "myq_gain_text": "C0392B",
    "iva_eval_bg": "E6F7F8", "iva_eval_text": "0E7C7B",
    "iva_gain_bg": "F0F4F8", "iva_gain_text": "2C5F7C",
}
FONT = "微软雅黑"

# ----------------------------- 底层工具 -----------------------------

def _set_cell_bg(cell, color_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), color_hex)
    tcPr.append(shd)

def _set_font(run, size=10.5, bold=False, color=None):
    run.font.name = FONT; run.font.size = Pt(size); run.font.bold = bold
    run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT)
    if color: run.font.color.rgb = RGBColor.from_string(color)

def add_para(doc_or_cell, text="", size=10.5, bold=False, color=None,
             align=None, space_after=6, space_before=0, indent=None):
    p = doc_or_cell.add_paragraph()
    if align: p.alignment = align
    pf = p.paragraph_format
    pf.space_after = Pt(space_after); pf.space_before = Pt(space_before); pf.line_spacing = 1.3
    if indent: pf.first_line_indent = Pt(indent)
    if text:
        _set_font(p.add_run(text), size, bold, color)
    return p

def _cell_text(cell, text, size=9.5, bold=False, color=None, align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ""
    for i, part in enumerate(str(text).split("\n")):
        p = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
        p.alignment = align
        p.paragraph_format.line_spacing = 1.3
        p.paragraph_format.space_after = Pt(2); p.paragraph_format.space_before = Pt(2)
        _set_font(p.add_run(part), size, bold, color)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP

def _page_break(doc):
    p = doc.add_paragraph()
    p.add_run().add_break()
    # python-docx add_break 默认是行内break，用 w:type page 确保分页
    for r in p.runs:
        for br in r._element.findall(qn('w:br')):
            br.set(qn('w:type'), 'page')

def _set_doc_defaults(doc):
    style = doc.styles['Normal']
    style.font.name = FONT; style.font.size = Pt(10.5)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), FONT)
    for section in doc.sections:
        section.top_margin = Cm(1.8); section.bottom_margin = Cm(1.8)
        section.left_margin = Cm(1.8); section.right_margin = Cm(1.8)

# ----------------------------- 封面 + 说明 -----------------------------

def build_cover_and_intro(doc, session):
    """首份文档才构建封面 + 使用说明。"""
    # 封面
    for _ in range(6): doc.add_paragraph()
    add_para(doc, "面试复盘录", size=32, bold=True, color=COLORS["primary"],
             align=WD_ALIGN_PARAGRAPH.CENTER, space_after=8)
    add_para(doc, f"{session.get('company','')} · {session.get('position','')}",
             size=18, color=COLORS["accent"], align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
    add_para(doc, "逐题原话复盘 · 持续沉淀", size=12, color=COLORS["muted"],
             align=WD_ALIGN_PARAGRAPH.CENTER, space_after=40)
    add_para(doc, f"候选人：{session.get('candidate','')}", size=12,
             align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
    add_para(doc, "（学校/院系）", size=12, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
    add_para(doc, f"首场记录：{session.get('date','')}", size=12, color=COLORS["muted"],
             align=WD_ALIGN_PARAGRAPH.CENTER)
    _page_break(doc)

    # 使用说明
    add_para(doc, "如何使用这份文档", size=16, bold=True, color=COLORS["primary"], space_after=10)
    add_para(doc, "这是一份可复用的面试复盘档案。每参加一场面试，录音转写后，用 interview-review skill 生成该场复盘，"
                  "并以 --append 模式追加到本文档末尾，即可累积成持续成长的复盘记录。", size=10.5, space_after=8)
    steps = [
        "1. 新建一级标题：「第 N 场面试 · 公司/岗位 · 日期」",
        "2. 补充该场的岗位JD与面试官业务线（便于复盘匹配度）",
        "3. 逐题填表格：序号/时间/面试官问题/我的回答/AI评价/改进版回答/提升点",
        "4. 末尾写「面试官反馈」和「核心复盘结论」",
        "5. 评价用符号：✅ 好 / 🟡 可改进 / 🔴 问题，便于一眼看出薄弱点",
    ]
    for s in steps:
        add_para(doc, s, size=10.5, space_after=4, indent=21)
    add_para(doc, "复盘的核心价值不在记录，而在「改进版回答」和「提升点」——"
                  "把每道没答好的题重写成下次能直接用的版本，并说清到底提升了什么。",
             size=10.5, bold=True, color=COLORS["gain_text"], space_before=8)
    _page_break(doc)

# ----------------------------- 业务线 -----------------------------

def build_business_line(doc, session):
    add_para(doc, f"第 {session.get('round',1)} 场面试 · {session.get('company','')}·{session.get('position','')} · {session.get('date','')}",
             size=18, bold=True, color=COLORS["primary"], space_after=4)
    result = session.get("result", "")
    if result:
        add_para(doc, f"结果：{result}", size=11, color=COLORS["gain_text"], space_after=12)

    add_para(doc, "一、岗位与业务线", size=14, bold=True, color=COLORS["primary"], space_after=6)

    jd = session.get("jd") or {}
    if jd.get("responsibilities"):
        add_para(doc, "【岗位 JD · 职责】", size=11, bold=True, space_after=4)
        for j in jd["responsibilities"]:
            add_para(doc, "• " + j, size=10, space_after=2, indent=21)
    if jd.get("requirements"):
        add_para(doc, "【职责要求】", size=11, bold=True, space_after=4, space_before=4)
        for j in jd["requirements"]:
            add_para(doc, "• " + j, size=10, space_after=2, indent=21)

    bl = session.get("business_line")
    if bl:
        add_para(doc, "【面试官所在业务线】", size=11, bold=True, space_after=4, space_before=8)
        add_para(doc, bl, size=10, space_after=4, indent=21)
    hint = session.get("match_hint")
    if hint:
        add_para(doc, "▶ " + hint, size=10, color=COLORS["gain_text"], space_after=8, indent=21, bold=True)
    _page_break(doc)

# ----------------------------- 逐题表格 -----------------------------

# 普通题行规格（面试官问、你答）
ROWS_NORMAL = [
    ("面试官问题", "q", "q_bg", "q_text", True),
    ("我的回答（原话）", "a", "a_bg", "a_text", False),
    ("AI 评价", "eval", "eval_bg", "eval_text", False),
    ("改进版回答", "better", "better_bg", "better_text", False),
    ("提升点（改进版 vs 原回答）", "gain", "gain_bg", "gain_text", False),
]

# 反问环节行规格（你问、面试官答，双侧评价）
ROWS_REVERSE = [
    ("我的提问（原话）", "q", "myq_bg", "myq_text", True),
    ("面试官回答（原话）", "a", "iva_bg", "iva_text", False),
    ("对提问的评价", "q_eval", "myq_eval_bg", "myq_eval_text", False),
    ("提问改进版", "q_better", "myq_better_bg", "myq_better_text", False),
    ("提问提升点", "q_gain", "myq_gain_bg", "myq_gain_text", False),
    ("对面试官回答的分析", "a_eval", "iva_eval_bg", "iva_eval_text", False),
    ("回答启示（这条反馈怎么用）", "a_gain", "iva_gain_bg", "iva_gain_text", False),
]

def _render_qa_title(doc, item, color_key="accent"):
    """渲染题目标题行：题号 〔时间〕· 模块"""
    title = f"{item.get('no','')}  〔{item.get('time','')}〕 · {item.get('module','')}"
    add_para(doc, title, size=11, bold=True, color=COLORS[color_key], space_before=12, space_after=4)

def _render_qa_table(doc, item, rows_spec):
    """渲染一道题的表格。rows_spec: [(label, key, bg_key, text_key, is_accent)]"""
    table = doc.add_table(rows=len(rows_spec), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    table.autofit = False
    for row in table.rows:
        row.cells[0].width = Cm(2.4); row.cells[1].width = Cm(15.0)

    for i, (label, key, bg_key, text_key, is_accent) in enumerate(rows_spec):
        c0, c1 = table.rows[i].cells
        _cell_text(c0, label, size=9.5, bold=True,
                   color="FFFFFF" if is_accent else COLORS["label_text"],
                   align=WD_ALIGN_PARAGRAPH.CENTER)
        _set_cell_bg(c0, COLORS[bg_key] if is_accent else COLORS["label_bg"])
        _cell_text(c1, item.get(key, ""), size=9.5, color=COLORS[text_key])
        _set_cell_bg(c1, COLORS[bg_key])

    # 行不跨页拆分
    for row in table.rows:
        trPr = row._tr.get_or_add_trPr()
        trPr.append(OxmlElement('w:cantSplit'))

def build_qa_section(doc, session):
    qa = session.get("qa", [])
    normal_items = [it for it in qa if it.get("type", "normal") != "reverse"]
    reverse_items = [it for it in qa if it.get("type") == "reverse"]

    # —— 二、逐题复盘（普通题）——
    if normal_items:
        add_para(doc, "二、逐题复盘（原话版）", size=14, bold=True, color=COLORS["primary"], space_after=4)
        add_para(doc, "说明：以下「面试官问题」与「我的回答」均为录音转写原话（已清理口吃重复，保留全部实质内容），未作概括缩写。",
                 size=9.5, color=COLORS["muted"], space_after=10)
        for item in normal_items:
            _render_qa_title(doc, item)
            _render_qa_table(doc, item, ROWS_NORMAL)

    # —— 三、反问环节（我的提问 vs 面试官回答）——
    if reverse_items:
        if normal_items:
            _page_break(doc)
        add_para(doc, "三、反问环节（我的提问 vs 面试官回答）", size=14, bold=True,
                 color=COLORS["primary"], space_after=4)
        add_para(doc, "说明：反问环节角色反转——「我的提问」要评价（问得好不好、得不得体），"
                      "「面试官回答」要分析（透露了什么信息/信号、对候选人有什么用）。",
                 size=9.5, color=COLORS["muted"], space_after=10)
        for item in reverse_items:
            _render_qa_title(doc, item, color_key="myq_text")
            _render_qa_table(doc, item, ROWS_REVERSE)

    _page_break(doc)

# ----------------------------- 反馈 + 结论 -----------------------------

def build_feedback_section(doc, session):
    # 有反问环节时章节顺移：反馈=四，结论=五；否则反馈=三，结论=四
    has_reverse = any(it.get("type") == "reverse" for it in session.get("qa", []))
    fb_num = "四" if has_reverse else "三"
    concl_num = "五" if has_reverse else "四"

    fb = session.get("interviewer_feedback")
    if fb:
        add_para(doc, f"{fb_num}、面试官原话反馈", size=14, bold=True, color=COLORS["primary"], space_after=6)
        add_para(doc, fb, size=10.5, space_after=8, indent=21)

    concls = session.get("conclusions") or []
    if concls:
        add_para(doc, f"{concl_num}、本场核心复盘结论", size=14, bold=True, color=COLORS["primary"], space_after=6, space_before=10)
        for item in concls:
            label, content = (item if isinstance(item, list) else (item.get("label",""), item.get("content","")))
            p = add_para(doc, "", space_after=4)
            _set_font(p.add_run(f"【{label}】"), 10.5, True, COLORS["gain_text"])
            _set_font(p.add_run(content), 10.5, False, COLORS["a_text"])

    add_para(doc, "", space_after=20)
    add_para(doc, "—— 下一场面试录音转写后，在此处之后追加新场次章节 ——",
             size=10, color=COLORS["muted"], align=WD_ALIGN_PARAGRAPH.CENTER, bold=True)

# ----------------------------- 主流程 -----------------------------

def build_session_section(doc, session):
    """把单个场次的所有章节写入 doc（append 模式也用这个）。"""
    build_business_line(doc, session)
    build_qa_section(doc, session)
    build_feedback_section(doc, session)

def main():
    ap = argparse.ArgumentParser(description="面试复盘 Word 文档生成器")
    ap.add_argument("--session", required=True, help="场次 JSON 数据文件")
    ap.add_argument("--out", required=True, help="输出 .docx 路径")
    ap.add_argument("--append", action="store_true", help="追加到已有文档")
    args = ap.parse_args()

    with open(args.session, "r", encoding="utf-8") as f:
        session = json.load(f)

    if args.append and os.path.exists(args.out):
        doc = Document(args.out)
    else:
        doc = Document()
        _set_doc_defaults(doc)
        build_cover_and_intro(doc, session)

    build_session_section(doc, session)
    doc.save(args.out)
    n_qa = len(session.get("qa", []))
    sys.stderr.write(f"[gen_review] wrote {args.out} ({'append' if args.append else 'new'}, {n_qa} Q&A)\n")
    return 0

if __name__ == "__main__":
    sys.exit(main())
